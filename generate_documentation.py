#!/usr/bin/env python3
"""
BachatKaro - Comprehensive Documentation Generator
Generates Professional PRD + Technical Documentation + Maintenance Guide
Output: BACHAT_KARO_PRD_TECHNICAL_DOCUMENTATION.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    heading.style = f'Heading {level}'
    return heading

def add_paragraph_styled(doc, text, bold=False, italic=False, color=None):
    """Add a paragraph with optional styling"""
    p = doc.add_paragraph(text)
    if bold or italic or color:
        run = p.runs[0] if p.runs else p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
    return p

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_table(doc, rows, cols, header_data=None, data_rows=None):
    """Create and format a table"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    if header_data:
        header_cells = table.rows[0].cells
        for i, header in enumerate(header_data):
            header_cells[i].text = header
            # Style header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(header_cells[i], '4472C4')
    
    if data_rows:
        for row_idx, row_data in enumerate(data_rows, start=1):
            if row_idx < len(table.rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        table.rows[row_idx].cells[col_idx].text = str(cell_data)
    
    return table

def generate_documentation():
    """Generate the complete documentation"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== COVER PAGE =====
    title = doc.add_heading('BACHAT KARO', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_heading('Intelligent Expense Tracker & Financial Planner for India', level=2)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc_type = doc.add_heading('COMBINED PRODUCT REQUIREMENTS DOCUMENT + TECHNICAL DOCUMENTATION', level=3)
    doc_type.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    version = doc.add_paragraph(f'Document Version: 1.0')
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    doc.add_heading('TABLE OF CONTENTS', level=1)
    
    toc_items = [
        'SECTION 1: EXECUTIVE SUMMARY',
        'SECTION 2: PRODUCT REQUIREMENTS DOCUMENT (PRD)',
        '  2.1 - Product Overview',
        '  2.2 - Target Users',
        '  2.3 - Features Overview',
        '  2.4 - User Journeys',
        '  2.5 - Functional Requirements',
        '  2.6 - Non-Functional Requirements',
        '  2.7 - UX & Design Guidelines',
        '  2.8 - Future Roadmap',
        'SECTION 3: TECHNICAL DOCUMENTATION',
        '  3.1 - System Overview',
        '  3.2 - Tech Stack',
        '  3.3 - Architecture & Data Flow',
        '  3.4 - Project Structure',
        '  3.5 - Module & Feature Breakdown',
        '  3.6 - Database & Models',
        '  3.7 - API Documentation',
        '  3.8 - Configuration',
        '  3.9 - Setup & Deployment',
        '  3.10 - Testing Strategy',
        'SECTION 4: MAINTENANCE & UPDATE GUIDE',
        '  4.1 - How to Modify Existing Features',
        '  4.2 - How to Add New Features',
        '  4.3 - Risk Areas & Critical Components',
        '  4.4 - Performance Optimization Opportunities',
        'SECTION 5: APPENDIX',
        '  5.1 - Glossary',
        '  5.2 - Assumptions & Known Issues',
        '  5.3 - Open Questions',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        if item.startswith('  '):
            p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # ===== SECTION 1: EXECUTIVE SUMMARY =====
    doc.add_heading('SECTION 1: EXECUTIVE SUMMARY', level=1)
    
    doc.add_heading('High-Level Overview', level=2)
    doc.add_paragraph(
        'BachatKaro is a modern, intelligent expense tracking and financial planning application '
        'designed specifically for Indian users. The platform empowers individuals and groups to manage '
        'expenses, split bills, track loans/EMIs, plan savings goals, and generate smart trip recommendations — '
        'all unified in a single intuitive interface.'
    )
    
    doc.add_heading('Current Maturity Level', level=2)
    doc.add_paragraph(
        'Status: Beta / Product Version 1.0'
    )
    doc.add_paragraph(
        'The application is feature-complete with core expense tracking, group management, EMI calculations, '
        'and trip planning. Infrastructure is stable with Supabase backend supporting real-time data synchronization. '
        'Mobile responsiveness is implemented; a dedicated mobile app is planned for future phases.'
    )
    
    doc.add_heading('Key Architecture Decisions', level=2)
    
    decisions = [
        ('Frontend Framework', 'React 18 with TypeScript for type safety and modern component architecture'),
        ('Build Tool', 'Vite for fast development and optimized production builds'),
        ('Database', 'Supabase (PostgreSQL) providing real-time updates, Row-Level Security, and serverless functions'),
        ('State Management', 'React Context API for auth/preferences + TanStack React Query for server state'),
        ('UI Library', 'shadcn-ui built on Radix UI for accessible, composable components'),
        ('Styling', 'Tailwind CSS with dark mode support via next-themes'),
        ('Voice Input', 'Web Speech API with language-specific keyword mapping for expense detection'),
        ('Deployment', 'Vercel for frontend; Supabase Cloud for backend'),
    ]
    
    for title, desc in decisions:
        p = doc.add_paragraph(f'{title}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Core Value Propositions', level=2)
    value_props = [
        'Voice-powered expense entry (Hindi/English/Hinglish) — fastest way to log expenses',
        'Intelligent group bill splitting with settlement tracking',
        'Comprehensive EMI/loan management with reducing & flat interest calculation',
        'AI-powered financial advice and wealth projections',
        'Smart trip planning with budget recommendations and local venue suggestions',
        'Bilingual interface supporting English and Hindi',
        'Real-time data synchronization across devices',
    ]
    for prop in value_props:
        doc.add_paragraph(prop, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECTION 2: PRODUCT REQUIREMENTS DOCUMENT =====
    doc.add_heading('SECTION 2: PRODUCT REQUIREMENTS DOCUMENT (PRD)', level=1)
    
    # 2.1 Product Overview
    doc.add_heading('2.1 Product Overview', level=2)
    
    doc.add_heading('Product Name', level=3)
    doc.add_paragraph('BachatKaro')
    doc.add_paragraph('"Bachat" = Savings, "Karo" = Do/Manage (Hindi origin)')
    
    doc.add_heading('Vision & Mission', level=3)
    doc.add_paragraph(
        'Vision: Empower every Indian to achieve financial wellness through intelligent, '
        'collaborative expense management.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Mission: Provide an intuitive, culturally-relevant platform that simplifies expense tracking, '
        'bill splitting, loan management, and savings planning — making financial control accessible to all.',
        style='List Bullet'
    )
    
    doc.add_heading('Problem Statement', level=3)
    problems = [
        'Manual expense tracking is tedious and error-prone',
        'Splitting bills among friends/roommates leads to confusion and disputes',
        'Loan/EMI management lacks clarity on outstanding balance and payment schedules',
        'No integrated tool for group trip budgeting and expense allocation',
        'Spending patterns are invisible; users lack savings recommendations',
        'Existing solutions lack Hindi support and Indian payment methods context',
    ]
    for problem in problems:
        doc.add_paragraph(problem, style='List Bullet')
    
    doc.add_heading('Goals & Success Metrics', level=3)
    goals = [
        ('User Acquisition', 'Onboard 10K+ active users in first 6 months'),
        ('Engagement', '> 80% users add expense at least 3x per week'),
        ('Voice Adoption', '> 40% of expense entries via voice input'),
        ('Group Features', '> 60% of users create or join at least one group'),
        ('Retention', 'Month-over-month retention rate > 85%'),
        ('Financial Literacy', 'Users report improved spending awareness after 1 month'),
    ]
    
    for goal_title, metric in goals:
        p = doc.add_paragraph(f'{goal_title}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(metric)
    
    doc.add_page_break()
    
    # 2.2 Target Users
    doc.add_heading('2.2 Target Users', level=2)
    
    doc.add_heading('User Personas', level=3)
    
    personas = [
        {
            'name': 'Priya (College Student)',
            'age': '22',
            'background': 'Computer Science student, shares apartment with 3 roommates',
            'pain_points': 'Complex household expense sharing, unclear who owes what',
            'goals': 'Track daily expenses, split rent/utilities automatically',
            'tech_comfort': 'High'
        },
        {
            'name': 'Rajesh (Young Professional)',
            'age': '28',
            'background': 'Software engineer, planning a wedding, has home loan',
            'pain_points': 'Juggling multiple EMIs, unclear savings trajectory',
            'goals': 'Manage EMIs, track savings toward wedding',
            'tech_comfort': 'High'
        },
        {
            'name': 'Meera (Parent & Homemaker)',
            'age': '35',
            'background': 'Manages household finances, plans family trips',
            'pain_points': 'Manual record keeping, no clear budget tracking',
            'goals': 'Simple expense logging, family budget management',
            'tech_comfort': 'Medium'
        },
        {
            'name': 'Arjun (Group Trip Organizer)',
            'age': '25',
            'background': 'Frequently organizes trips with friends',
            'pain_points': 'Manual bill tracking during trips, complex settlement calculations',
            'goals': 'Quick expense splitting, automatic settlement suggestions',
            'tech_comfort': 'High'
        },
    ]
    
    for persona in personas:
        p = doc.add_paragraph(f"{persona['name']} ({persona['age']} years old)", style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        
        doc.add_paragraph(f"Background: {persona['background']}", style='List Bullet 2')
        doc.add_paragraph(f"Pain Points: {persona['pain_points']}", style='List Bullet 2')
        doc.add_paragraph(f"Goals: {persona['goals']}", style='List Bullet 2')
        doc.add_paragraph(f"Tech Comfort: {persona['tech_comfort']}", style='List Bullet 2')
    
    doc.add_heading('User Needs', level=3)
    needs = [
        'Simple, voice-powered expense entry (hands-free)',
        'Clear breakdown of where money is being spent',
        'Fair and automatic bill splitting with friends',
        'Transparent view of who owes whom',
        'Clear understanding of loan payoff schedule',
        'Ability to set and track savings goals',
        'Smart recommendations for financial improvement',
        'Support for Hindi language interface',
    ]
    for need in needs:
        doc.add_paragraph(need, style='List Bullet')
    
    doc.add_page_break()
    
    # 2.3 Features Overview
    doc.add_heading('2.3 Features Overview', level=2)
    
    features_data = [
        ['Feature', 'Description', 'Priority', 'Status'],
        ['Personal Expense Tracking', 'Log expenses manually or via voice (Eng/Hindi)', 'P0', 'Live'],
        ['Voice Input Support', 'Voice-to-text with automatic category detection', 'P0', 'Live'],
        ['Dashboard Analytics', 'Charts, trends, spending breakdown by category', 'P0', 'Live'],
        ['Group Creation & Management', 'Create groups, invite members, manage permissions', 'P0', 'Live'],
        ['Bill Splitting', 'Split expenses equally or custom ratios', 'P0', 'Live'],
        ['Settlement Tracking', 'Track who owes whom with visual indicators', 'P0', 'Live'],
        ['EMI/Loan Management', 'Add loans, calculate EMI, track remaining balance', 'P0', 'Live'],
        ['Savings Goals Tracking', 'Create goals, monitor progress with visual indicators', 'P1', 'Live'],
        ['Financial Health Score', 'Automated score based on spending habits', 'P1', 'Live'],
        ['Wealth Prediction', 'Project future savings based on current patterns', 'P1', 'Live'],
        ['Trip Planning Assistant', 'Generate trip itineraries with budget & recommendations', 'P2', 'Live'],
        ['Bill Roulette Game', 'Fun game to randomly select who pays', 'P2', 'Live'],
        ['WhatsApp Integration', 'Share expenses, invoices, and trips via WhatsApp', 'P2', 'Live'],
        ['Dark Mode', 'Full dark mode support across all pages', 'P3', 'Live'],
        ['Mobile Responsiveness', 'Fully responsive design for mobile/tablet', 'P1', 'Live'],
        ['Export Data', 'Export expenses as CSV/PDF', 'P3', 'Planned'],
        ['Recurring Expenses', 'Automated recurring expense logging', 'P2', 'Planned'],
        ['Multi-currency Support', 'Support for currencies beyond INR (future)', 'P3', 'Backlog'],
        ['Offline Mode', 'Limited functionality without internet', 'P3', 'Backlog'],
    ]
    
    create_table(doc, len(features_data), len(features_data[0]), 
                header_data=features_data[0], data_rows=features_data[1:])
    
    doc.add_page_break()
    
    # 2.4 User Journeys
    doc.add_heading('2.4 User Journeys', level=2)
    
    doc.add_heading('Journey 1: Onboarding & First Expense', level=3)
    steps = [
        'User visits BachatKaro.com',
        'Clicks "Sign Up", enters email & password',
        'Account created; auth email sent',
        'Redirected to Setup Wizard (country selection)',
        'Completes setup, lands on Dashboard',
        'Clicks voice mic button 🎤',
        'Says "500 for coffee"',
        'App automatically categorizes as Food, confirms "₹500 Added ✅"',
        'Expense visible in Recent Expenses list',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    doc.add_heading('Journey 2: Splitting a Restaurant Bill', level=3)
    steps2 = [
        'User goes to "Groups" tab',
        'Clicks "Create Group" — names it "Goa Trip Friends"',
        'Adds 3 friends names',
        'Sends group invite code to friends (they join)',
        'Group admin clicks "Add Expense", enters ₹2000 restaurant bill',
        'Selects "All members split equally"',
        'App calculates: each person owes ₹500',
        'Dashboard shows "Settlements": Person A owes ₹500, Person B owes ₹500, etc.',
        'Person A clicks "Settle" → payment link for UPI',
        'All marked as settled ✅',
    ]
    for i, step in enumerate(steps2, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    doc.add_heading('Journey 3: EMI Loan Tracking', level=3)
    steps3 = [
        'User goes to "Planning" tab',
        'Clicks "Add EMI" → enters car loan details (₹5,00,000 @ 9% for 60 months)',
        'Selects "Reducing Balance" interest method',
        'App calculates: ₹9,655 monthly EMI',
        'EMI card shows outstanding balance, interest paid so far',
        'User can see full amortization schedule',
        'App alerts when EMI payment due date approaches',
    ]
    for i, step in enumerate(steps3, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    doc.add_page_break()
    
    # 2.5 Functional Requirements
    doc.add_heading('2.5 Functional Requirements', level=2)
    
    doc.add_heading('Expense Management', level=3)
    fr_expense = [
        'Users can add expenses via manual form OR voice input',
        'Voice input auto-detects category from keywords (Food, Travel, Shopping, Bills)',
        'Expenses include: amount, category, payment mode (UPI/Cash/Card/Net Banking), date, notes',
        'Users can edit/delete their expenses anytime',
        'Expenses are real-time synced across devices',
        'Support for multiple payment modes per expense',
    ]
    for req in fr_expense:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('Group Expense & Bill Splitting', level=3)
    fr_group = [
        'Users can create groups with custom names',
        'Groups have admin (creator) and members',
        'Members can be added by email or invite code',
        'Group expenses can be split equally or custom (unequal)',
        'System calculates net balance for each user in group (who owes what)',
        'Settlement can be marked manually or via payment gateway integration',
    ]
    for req in fr_group:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('EMI/Loan Management', level=3)
    fr_emi = [
        'Users can add multiple loans (car, home, personal, etc.)',
        'Each loan tracks: principal, annual rate, tenure (months), start date',
        'System supports Reducing Balance & Flat Interest methods',
        'App calculates: monthly EMI, total interest, outstanding balance, amortization schedule',
        'EMI details include principal component, interest component for current month',
        'Users can track loan progress visually',
    ]
    for req in fr_emi:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('Analytics & Insights', level=3)
    fr_analytics = [
        'Dashboard shows spending by category (pie/bar chart)',
        'Trend chart shows spending over time (daily, weekly, monthly)',
        'Financial Health Score (0-100) based on: savings rate, expense control, diversification',
        'Smart Financial Mentor provides AI tips based on patterns',
        'Wealth Predictor projects balance over 6/12 months',
    ]
    for req in fr_analytics:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_page_break()
    
    # 2.6 Non-Functional Requirements
    doc.add_heading('2.6 Non-Functional Requirements', level=2)
    
    doc.add_heading('Performance', level=3)
    doc.add_paragraph('Page Load Time: < 2 seconds on 4G connection', style='List Bullet')
    doc.add_paragraph('API Response: < 500ms for expense queries', style='List Bullet')
    doc.add_paragraph('Voice recognition: Response within 1-2 seconds of saying expense', style='List Bullet')
    doc.add_paragraph('Real-time sync: Data synced within 500ms across devices', style='List Bullet')
    
    doc.add_heading('Security', level=3)
    doc.add_paragraph('All data encrypted in transit (TLS 1.2+)', style='List Bullet')
    doc.add_paragraph('User passwords hashed using bcrypt', style='List Bullet')
    doc.add_paragraph('Row-Level Security (RLS) enforced at database level', style='List Bullet')
    doc.add_paragraph('Users can only access their own data & groups they belong to', style='List Bullet')
    doc.add_paragraph('Sensitive data (PII) not logged; audit trails maintained', style='List Bullet')
    
    doc.add_heading('Scalability', level=3)
    doc.add_paragraph('Database: PostgreSQL with auto-scaling (Supabase)', style='List Bullet')
    doc.add_paragraph('Frontend: Serverless deployment on Vercel with edge functions', style='List Bullet')
    doc.add_paragraph('Real-time: PostgreSQL LISTEN/NOTIFY for live updates', style='List Bullet')
    doc.add_paragraph('Target: Support 100K+ concurrent users', style='List Bullet')
    
    doc.add_heading('Accessibility', level=3)
    doc.add_paragraph('WCAG 2.1 AA compliance', style='List Bullet')
    doc.add_paragraph('Keyboard navigation support', style='List Bullet')
    doc.add_paragraph('Screen reader compatible (ARIA labels)', style='List Bullet')
    doc.add_paragraph('Bilingual support (English & Hindi)', style='List Bullet')
    doc.add_paragraph('Dark mode for user comfort', style='List Bullet')
    
    doc.add_page_break()
    
    # 2.7 UX & Design Guidelines
    doc.add_heading('2.7 UX & Design Guidelines', level=2)
    
    doc.add_heading('Navigation Structure', level=3)
    nav_items = [
        'Home/Dashboard: Primary entry point with 5 main tabs',
        'Daily Tracker: Expense viewing & quick-add',
        'Planning: Salary, budget, EMI management',
        'Future Wealth: Projections and advisor',
        'My Dreams: Savings goals tracking',
        'Groups: Group management & bill splitting',
        'Top Nav: App logo, search (future), user menu',
        'Bottom Nav (Mobile): Sticky tabs for easy access',
    ]
    for nav in nav_items:
        doc.add_paragraph(nav, style='List Bullet')
    
    doc.add_heading('Layout Rules', level=3)
    doc.add_paragraph('Max content width: 1200px (desktop)', style='List Bullet')
    doc.add_paragraph('Mobile-first responsive design: 320px, 768px, 1024px breakpoints', style='List Bullet')
    doc.add_paragraph('Cards: Use shadcn Card component with subtle shadows', style='List Bullet')
    doc.add_paragraph('Colors: Indigo primary (#4f46e5), Purple accent (#9333ea)', style='List Bullet')
    doc.add_paragraph('Spacing: 8px base unit (8, 16, 24, 32, etc.)', style='List Bullet')
    
    doc.add_heading('UI Consistency Rules', level=3)
    doc.add_paragraph('All buttons: Use shadcn Button component', style='List Bullet')
    doc.add_paragraph('Forms: Use shadcn form inputs with label above', style='List Bullet')
    doc.add_paragraph('Modals: Use shadcn Dialog for all modal interactions', style='List Bullet')
    doc.add_paragraph('Toasts: Use Sonner for notifications (success/error/warning)', style='List Bullet')
    doc.add_paragraph('Icons: Lucide React for all icons (consistent icon set)', style='List Bullet')
    doc.add_paragraph('Charts: Recharts for all data visualizations', style='List Bullet')
    
    doc.add_page_break()
    
    # 2.8 Future Roadmap
    doc.add_heading('2.8 Future Roadmap', level=2)
    
    doc.add_heading('Q2 2026: Enhanced Collaboration', level=3)
    doc.add_paragraph('Recurring expense automation', style='List Bullet')
    doc.add_paragraph('Export to CSV/PDF', style='List Bullet')
    doc.add_paragraph('Scheduled reminders for bill pay dates', style='List Bullet')
    doc.add_paragraph('Budget alerts (over-spending notifications)', style='List Bullet')
    
    doc.add_heading('Q3 2026: Mobile App & Payments', level=3)
    doc.add_paragraph('Native iOS/Android app', style='List Bullet')
    doc.add_paragraph('Direct payment integration (Razorpay/PhonePe)', style='List Bullet')
    doc.add_paragraph('Rent/insurance bill splitting templates', style='List Bullet')
    
    doc.add_heading('Q4 2026 & Beyond', level=3)
    doc.add_paragraph('Multi-currency support (USD, EUR, etc.)', style='List Bullet')
    doc.add_paragraph('Investment tracking (stocks, mutual funds)', style='List Bullet')
    doc.add_paragraph('Integration with bank APIs for auto-categorization', style='List Bullet')
    doc.add_paragraph('AI-powered chatbots for natural language queries', style='List Bullet')
    doc.add_paragraph('Offline mode with sync when online', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECTION 3: TECHNICAL DOCUMENTATION =====
    doc.add_heading('SECTION 3: TECHNICAL DOCUMENTATION', level=1)
    
    # 3.1 System Overview
    doc.add_heading('3.1 System Overview', level=2)
    
    doc.add_heading('Architecture Diagram (Conceptual)', level=3)
    doc.add_paragraph(
        'User Browser (React App) ←→ Vite Dev Server / Vercel (Production) ←→ Supabase API ←→ PostgreSQL Database'
    )
    
    doc.add_heading('System Components', level=3)
    components = [
        ('Frontend', 'React 18 + TypeScript application running in browser'),
        ('Build Tool', 'Vite for development and production optimization'),
        ('State Management', 'React Context (Auth, Language) + React Query (server state)'),
        ('Backend', 'Supabase providing PostgreSQL database + auth + realtime + serverless functions'),
        ('Real-time Sync', 'PostgreSQL LISTEN/NOTIFY via Supabase Realtime'),
        ('Authentication', 'Supabase Auth (email/password + password reset)'),
        ('Deployment', 'Vercel (frontend), Supabase Cloud (backend)'),
    ]
    
    for comp_name, comp_desc in components:
        p = doc.add_paragraph(f'{comp_name}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(comp_desc)
    
    doc.add_page_break()
    
    # 3.2 Tech Stack
    doc.add_heading('3.2 Tech Stack', level=2)
    
    doc.add_heading('Frontend Stack', level=3)
    frontend = [
        ('Framework', 'React 18.3.1 - Component-based UI library'),
        ('Language', 'TypeScript 5.8 - Type-safe JavaScript'),
        ('Build Tool', 'Vite 5.4 - Fast development & production builds'),
        ('Router', 'React Router v6.30 - Client-side routing'),
        ('State Query', 'TanStack React Query 5.83 - Server state management'),
        ('Forms', 'React Hook Form 7.61 + Zod 3.25 - Form handling & validation'),
        ('Component Library', 'shadcn-ui with Radix UI - Accessible composable components'),
        ('Styling', 'Tailwind CSS 3.4 - Utility-first CSS framework'),
        ('Dark Mode', 'next-themes 0.3 - Theme switching'),
        ('Charts', 'Recharts 2.15 - React charting library'),
        ('Icons', 'Lucide React 0.462 - SVG icon library'),
        ('Toast/Alerts', 'Sonner 1.7.4 - Toast notifications'),
        ('UI Dialogs', 'Radix Dialog via shadcn - Modal dialogs'),
    ]
    
    for tech, desc in frontend:
        p = doc.add_paragraph(f'{tech}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Backend Stack', level=3)
    backend = [
        ('Database', 'PostgreSQL 15+ via Supabase'),
        ('Authentication', 'Supabase Auth with email/password'),
        ('Real-time', 'Supabase Realtime (PostgreSQL LISTEN/NOTIFY)'),
        ('API Layer', 'Supabase REST API + RPC (stored procedures)'),
        ('Serverless Functions', 'Supabase Edge Functions (Node.js/Deno) - email sending'),
        ('File Storage', 'Supabase Storage (future for receipts)'),
        ('Security', 'Row-Level Security (RLS) policies'),
    ]
    
    for tech, desc in backend:
        p = doc.add_paragraph(f'{tech}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Development Tools', level=3)
    dev_tools = [
        ('Testing', 'Vitest 3.2.4 - Unit testing framework'),
        ('Testing Library', 'Testing Library React 16 - Component testing'),
        ('Linting', 'ESLint 9.32 - JavaScript/TypeScript linting'),
        ('Package Manager', 'Bun - Fast JavaScript package manager'),
        ('IDE', 'VS Code recommended'),
    ]
    
    for tool, desc in dev_tools:
        p = doc.add_paragraph(f'{tool}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 3.3 Architecture & Data Flow
    doc.add_heading('3.3 Architecture & Data Flow', level=2)
    
    doc.add_heading('State Management Architecture', level=3)
    doc.add_paragraph('AuthContext: Manages user authentication state (login, logout, session)', style='List Bullet')
    doc.add_paragraph('LanguageContext: Manages UI language (English/Hindi) and user preferences', style='List Bullet')
    doc.add_paragraph('ThemeContext (via next-themes): Manages dark/light mode', style='List Bullet')
    doc.add_paragraph('React Query: Caches server state (expenses, groups, EMIs, goals)', style='List Bullet')
    
    doc.add_heading('Data Flow Example: Adding Expense', level=3)
    doc.add_paragraph('1. User speaks "500 for coffee" into voice input', style='List Number')
    doc.add_paragraph('2. VoiceExpenseAdder captures audio via Web Speech API', style='List Number')
    doc.add_paragraph('3. voiceParser.ts parses transcript → extracts amount, category (Food)', style='List Number')
    doc.add_paragraph('4. User confirms/edits details in form', style='List Number')
    doc.add_paragraph('5. onSubmit() calls supabase.from("expenses").insert({...})', style='List Number')
    doc.add_paragraph('6. Database INSERT triggers real-time broadcast via Supabase Realtime', style='List Number')
    doc.add_paragraph('7. React Query invalidates expenses cache, refetches updated list', style='List Number')
    doc.add_paragraph('8. Dashboard components re-render with new expense visible', style='List Number')
    doc.add_paragraph('9. All other user devices receive real-time update', style='List Number')
    
    doc.add_heading('Real-time Synchronization', level=3)
    doc.add_paragraph(
        'Supabase Realtime uses PostgreSQL\'s LISTEN/NOTIFY to push changes to connected clients. '
        'When an expense is added, the database broadcast immediately updates all subscribed clients. '
        'React Query automatically refetches and updates the UI.',
        style='List Bullet'
    )
    
    doc.add_page_break()
    
    # 3.4 Project Structure
    doc.add_heading('3.4 Project Structure', level=2)
    
    project_structure = """
src/
├── App.tsx                    # Main app router & layout provider
├── main.tsx                   # React entry point
├── index.css                  # Global styles
├── App.css                    # App-specific styles
├── pages/                     # Page/route components
│   ├── Dashboard.tsx          # Main expense tracking dashboard
│   ├── AddExpense.tsx         # Dedicated expense add page
│   ├── GroupExpenses.tsx      # Group management & bill splitting
│   ├── Savings.tsx            # Savings goals tracking
│   ├── Analytics.tsx          # Lazy-loaded analytics page
│   ├── Auth.tsx               # Login/registration page
│   ├── TripPlanView.tsx       # Trip plan display
│   ├── TripShareHandler.tsx   # Deep link handler for shared trips
│   ├── SetupWizard.tsx        # Onboarding wizard
│   ├── JoinGroup.tsx          # Join group via code
│   └── NotFound.tsx           # 404 page
│
├── components/
│   ├── auth/                  # Authentication components
│   │   ├── LoginForm.tsx
│   │   ├── RegisterForm.tsx
│   │   └── ForgotPassword.tsx
│   │
│   ├── dashboard/             # Dashboard-specific components
│   │   ├── ExpenseTotalsGrid.tsx
│   │   ├── RecentExpenses.tsx
│   │   ├── CategoryChart.tsx
│   │   ├── DateTrendChart.tsx
│   │   ├── SmartFinancialMentor.tsx
│   │   ├── FinancialHealthScore.tsx
│   │   ├── FutureWealthPredictor.tsx
│   │   ├── EditExpenseDialog.tsx
│   │   ├── DeleteExpenseDialog.tsx
│   │   └── [... other dashboard components]
│   │
│   ├── groups/                # Group-specific components
│   │   ├── BillRoulette.tsx
│   │   ├── TripAdvisor.tsx
│   │   └── tripProvider.ts
│   │
│   ├── layout/                # Layout components
│   │   ├── AppHeader.tsx
│   │   ├── DashboardSubheader.tsx
│   │   └── ProtectedRoute.tsx
│   │
│   ├── ui/                    # shadcn-ui components
│   │   ├── button.tsx
│   │   ├── card.tsx
│   │   ├── input.tsx
│   │   ├── dialog.tsx
│   │   ├── select.tsx
│   │   ├── tabs.tsx
│   │   ├── [... 20+ other shadcn components]
│   │   └── VoiceInput.tsx
│   │
│   ├── EMIBillsCard.tsx
│   ├── EMILoanDetailsBlock.tsx
│   ├── MonthlySnapshot.tsx
│   └── NavLink.tsx
│
├── contexts/                  # React Context providers
│   ├── AuthContext.tsx        # Auth state (user, session)
│   ├── LanguageContext.tsx    # i18n state (language, translations)
│   └── ThemeContext.tsx       # Dark mode context
│
├── hooks/                     # Custom React hooks
│   ├── useOptimizedList.ts
│   ├── use-mobile.tsx         # Mobile detection hook
│   └── use-toast.ts           # Toast notification hook
│
├── services/                  # API & business logic
│   ├── api.ts                 # Supabase function calls
│   ├── tripPlanner.ts         # Trip planning logic
│   ├── tripShareService.ts    # Trip sharing URLs
│   ├── deepLinkHandler.ts     # Deep link routing
│   └── mentorService.ts       # AI advisor logic
│
├── utils/                     # Helper utilities
│   ├── voiceParser.ts         # Voice transcript parsing
│   ├── loanCalculator.ts      # EMI/loan calculations
│   ├── currencyFormatter.ts   # Currency formatting
│   ├── transactionParser.ts   # Transaction parsing
│   ├── validators.ts          # Form validators
│   ├── loanCalculator.ts      # Loan math
│   ├── destinationGradient.ts # Trip destination styling
│   └── whatsappSummary.ts     # WhatsApp message templates
│
├── features/                  # Feature-specific modules
│   ├── analytics/             # Analytics-related features
│   └── split-expense/         # Expense splitting logic
│       ├── utils/
│       │   ├── splitCalculator.ts
│       │   ├── simplifyDebts.ts
│       │   └── [... other utilities]
│       └── types/
│
├── types/                     # TypeScript type definitions
│   ├── emi.ts                 # EMI-related types
│   ├── analytics.ts           # Analytics types
│   └── [... other types]
│
├── data/                      # Static data files
│   ├── DestinationsData.ts    # Trip destination database
│   └── tripData.ts            # Trip planning data
│
├── i18n/                      # Internationalization
│   ├── index.ts               # i18n configuration
│   └── translations.ts        # Language strings (EN & HI)
│
├── integrations/
│   └── supabase/
│       └── client.ts          # Supabase client initialization
│
├── config/                    # Configuration
│   ├── appConfig.ts           # App-level settings
│   ├── currency.ts            # Currency configuration
│   └── components.json        # shadcn-ui config
│
├── styles/
│   ├── globals.css            # Global CSS
│   └── [... component styles]
│
└── test/
    ├── setup.ts               # Vitest setup
    └── example.test.ts        # Example test
"""
    
    doc.add_paragraph(project_structure, style='List Bullet')
    
    doc.add_page_break()
    
    # 3.5 Module & Feature Breakdown
    doc.add_heading('3.5 Module & Feature Breakdown', level=2)
    
    modules = [
        {
            'name': 'Authentication Module',
            'purpose': 'Handle user signup, login, password reset, session management',
            'files': ['Auth.tsx', 'AuthContext.tsx', 'LoginForm.tsx', 'RegisterForm.tsx', 'ForgotPassword.tsx'],
            'key_functions': ['signUp()', 'signIn()', 'resetPassword()', 'useAuth()'],
            'dependencies': ['Supabase Auth', 'React Router'],
        },
        {
            'name': 'Expense Tracking Module',
            'purpose': 'Add, edit, delete, view expenses with voice support',
            'files': ['Dashboard.tsx', 'AddExpense.tsx', 'voiceParser.ts', 'VoiceInput.tsx', 'RecentExpenses.tsx'],
            'key_functions': ['parseVoiceTranscript()', 'addExpense()', 'editExpense()', 'deleteExpense()'],
            'dependencies': ['Web Speech API', 'Supabase REST', 'React Query'],
        },
        {
            'name': 'Group Expense & Bill Splitting',
            'purpose': 'Manage group expenses, split bills, calculate settlements',
            'files': ['GroupExpenses.tsx', 'splitCalculator.ts', 'simplifyDebts.ts', 'BillRoulette.tsx'],
            'key_functions': ['calculateSplit()', 'computeBalances()', 'simplifyDebts()'],
            'dependencies': ['Supabase', 'React Query'],
        },
        {
            'name': 'EMI/Loan Management',
            'purpose': 'Track loans, calculate EMI, manage payment schedule',
            'files': ['Dashboard.tsx', 'EMILoanDetailsBlock.tsx', 'loanCalculator.ts'],
            'key_functions': ['getLoanSummary()', 'calculateReducingEMI()', 'getMonthsPaid()'],
            'dependencies': ['Date-fns'],
        },
        {
            'name': 'Analytics & Insights',
            'purpose': 'Display spending trends, financial health, wealth projections',
            'files': ['Analytics.tsx', 'CategoryChart.tsx', 'FutureWealthPredictor.tsx', 'SmartFinancialMentor.tsx'],
            'key_functions': ['calculateHealthScore()', 'predictWealth()', 'generateAdvice()'],
            'dependencies': ['Recharts', 'Supabase'],
        },
        {
            'name': 'Trip Planning & Sharing',
            'purpose': 'Generate trip itineraries, create shareable links',
            'files': ['TripAdvisor.tsx', 'tripPlanner.ts', 'tripShareService.ts', 'TripPlanView.tsx'],
            'key_functions': ['generatePlan()', 'createShareableLink()', 'parseDeepLink()'],
            'dependencies': ['Destination Data', 'React Query'],
        },
    ]
    
    for mod in modules:
        doc.add_heading(f"{mod['name']}", level=3)
        doc.add_paragraph(f"Purpose: {mod['purpose']}", style='List Bullet')
        doc.add_paragraph(f"Files: {', '.join(mod['files'])}", style='List Bullet')
        doc.add_paragraph(f"Key Functions: {', '.join(mod['key_functions'])}", style='List Bullet')
        doc.add_paragraph(f"Dependencies: {', '.join(mod['dependencies'])}", style='List Bullet')
    
    doc.add_page_break()
    
    # 3.6 Database & Models
    doc.add_heading('3.6 Database & Models', level=2)
    
    doc.add_heading('Core Tables', level=3)
    
    tables_data = [
        ['Table Name', 'Purpose', 'Key Columns', 'Relationships'],
        ['users', 'Supabase auth users (auto-created by Auth)', 'id, email, created_at', 'Referenced in all user-owned tables'],
        ['user_preferences', 'User language, country settings', 'user_id, language, country', 'One-to-one with users'],
        ['expenses', 'Personal expense records', 'id, user_id, amount, category, date, payment_mode', 'user_id → users.id'],
        ['savings_goals', 'User savings targets', 'id, user_id, title, target_amount, current_amount, deadline', 'user_id → users.id'],
        ['emis', 'Loan/EMI records', 'id, user_id, name, principal, rate, tenure_months, start_date', 'user_id → users.id'],
        ['groups', 'Shared expense groups', 'id, created_by, name, created_at', 'created_by → users.id'],
        ['group_members', 'Group membership', 'group_id, user_id, joined_at', 'group_id → groups.id, user_id → users.id'],
        ['group_expenses', 'Expenses within groups', 'id, group_id, paid_by, amount, title, notes, created_at', 'group_id → groups.id, paid_by → users.id'],
        ['group_expense_participants', 'Who splits an expense', 'group_expense_id, user_id, amount_owed', 'group_expense_id → group_expenses.id'],
    ]
    
    create_table(doc, len(tables_data), len(tables_data[0]), 
                header_data=tables_data[0], data_rows=tables_data[1:])
    
    doc.add_heading('Table Schemas (DDL)', level=3)
    doc.add_paragraph('See supabase/migrations/ folder for SQL definitions. Key examples:', style='List Bullet')
    
    doc.add_paragraph('expenses:', style='List Number')
    doc.add_paragraph(
        'CREATE TABLE expenses (\n'
        '  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
        '  user_id UUID NOT NULL REFERENCES auth.users(id),\n'
        '  amount DECIMAL(10, 2) NOT NULL,\n'
        '  category TEXT NOT NULL,\n'
        '  payment_mode TEXT,\n'
        '  date DATE NOT NULL,\n'
        '  description TEXT,\n'
        '  created_at TIMESTAMP DEFAULT NOW(),\n'
        '  updated_at TIMESTAMP DEFAULT NOW()\n'
        ');',
        style='List Bullet 2'
    )
    
    doc.add_paragraph('groups:', style='List Number')
    doc.add_paragraph(
        'CREATE TABLE groups (\n'
        '  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
        '  created_by UUID NOT NULL REFERENCES auth.users(id),\n'
        '  name TEXT NOT NULL,\n'
        '  created_at TIMESTAMP DEFAULT NOW()\n'
        ');',
        style='List Bullet 2'
    )
    
    doc.add_heading('Row-Level Security (RLS) Policies', level=3)
    doc.add_paragraph(
        'All tables have RLS enabled. Users can only read/write their own data. '
        'Example: Users can only access expenses WHERE user_id = auth.uid()',
        style='List Bullet'
    )
    
    doc.add_page_break()
    
    # 3.7 API Documentation
    doc.add_heading('3.7 API Documentation', level=2)
    
    doc.add_heading('Authentication Endpoints', level=3)
    
    api_endpoints = [
        {
            'name': 'Sign Up',
            'method': 'POST',
            'endpoint': '/auth/v1/signup',
            'provider': 'Supabase',
            'input': '{ email, password }',
            'output': '{ user, session }',
            'notes': 'Creates new user, sends verification email'
        },
        {
            'name': 'Sign In',
            'method': 'POST',
            'endpoint': '/auth/v1/token?grant_type=password',
            'provider': 'Supabase',
            'input': '{ email, password }',
            'output': '{ access_token, refresh_token, user }',
            'notes': 'Returns JWT token valid for 1 hour'
        },
        {
            'name': 'Reset Password',
            'method': 'POST',
            'endpoint': '/auth/v1/recovery',
            'provider': 'Supabase',
            'input': '{ email }',
            'output': '{ success: true }',
            'notes': 'Sends reset link to email'
        },
    ]
    
    for endpoint in api_endpoints:
        doc.add_paragraph(f"{endpoint['name']}", style='List Number')
        doc.add_paragraph(f"Method: {endpoint['method']}", style='List Bullet 2')
        doc.add_paragraph(f"Endpoint: {endpoint['endpoint']}", style='List Bullet 2')
        doc.add_paragraph(f"Input: {endpoint['input']}", style='List Bullet 2')
        doc.add_paragraph(f"Output: {endpoint['output']}", style='List Bullet 2')
        doc.add_paragraph(f"Notes: {endpoint['notes']}", style='List Bullet 2')
    
    doc.add_heading('Expense Endpoints (REST + RPC)', level=3)
    
    doc.add_paragraph('GET /rest/v1/expenses', style='List Number')
    doc.add_paragraph('Fetch user\'s expenses. Query: ?user_id=eq.{userId}', style='List Bullet 2')
    doc.add_paragraph('Returns: [{ id, user_id, amount, category, date, ... }]', style='List Bullet 2')
    
    doc.add_paragraph('POST /rest/v1/expenses', style='List Number')
    doc.add_paragraph('Create expense. Body: { user_id, amount, category, date, ... }', style='List Bullet 2')
    doc.add_paragraph('Returns: { id, created_at, ... }', style='List Bullet 2')
    
    doc.add_paragraph('PATCH /rest/v1/expenses?id=eq.{expenseId}', style='List Number')
    doc.add_paragraph('Update expense. Body: { amount, category, ... }', style='List Bullet 2')
    
    doc.add_paragraph('DELETE /rest/v1/expenses?id=eq.{expenseId}', style='List Number')
    doc.add_paragraph('Delete expense', style='List Bullet 2')
    
    doc.add_heading('Group Endpoints', level=3)
    
    doc.add_paragraph('GET /rest/v1/groups', style='List Number')
    doc.add_paragraph('Fetch user\'s groups (via group_members join)', style='List Bullet 2')
    
    doc.add_paragraph('POST /rest/v1/groups', style='List Number')
    doc.add_paragraph('Create new group. Body: { created_by, name }', style='List Bullet 2')
    
    doc.add_paragraph('POST /rest/v1/group_members', style='List Number')
    doc.add_paragraph('Add member to group. Body: { group_id, user_id }', style='List Bullet 2')
    
    doc.add_heading('RPC Functions', level=3)
    
    doc.add_paragraph('get_or_create_user_preferences(p_user_id)', style='List Number')
    doc.add_paragraph('Returns or creates user language/country preference', style='List Bullet 2')
    
    doc.add_paragraph('get_group_settlements(p_group_id)', style='List Number')
    doc.add_paragraph('Returns settlement calculations for a group (who owes whom)', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 3.8 Configuration
    doc.add_heading('3.8 Configuration', level=2)
    
    doc.add_heading('Environment Variables', level=3)
    doc.add_paragraph(
        'Create a `.env.local` file in project root with:'
    )
    
    env_vars = [
        'VITE_SUPABASE_URL=https://[your-project].supabase.co',
        'VITE_SUPABASE_ANON_KEY=eyJhbGciOiJIUzI1NiIsInR5cCI...',
        'VITE_ENABLE_ANALYTICS=true',
        'VITE_APP_NAME=BachatKaro',
    ]
    
    for env_var in env_vars:
        doc.add_paragraph(env_var, style='List Bullet')
    
    doc.add_heading('Configuration Files', level=3)
    doc.add_paragraph('vite.config.ts: Build configuration', style='List Bullet')
    doc.add_paragraph('tsconfig.json: TypeScript configuration', style='List Bullet')
    doc.add_paragraph('tailwind.config.ts: Tailwind CSS customization', style='List Bullet')
    doc.add_paragraph('postcss.config.js: PostCSS plugins', style='List Bullet')
    doc.add_paragraph('eslint.config.js: Linting rules', style='List Bullet')
    doc.add_paragraph('components.json: shadcn-ui configuration', style='List Bullet')
    
    doc.add_heading('App Configuration (appConfig.ts)', level=3)
    doc.add_paragraph('defaultCountry: "IN" (India)', style='List Bullet')
    doc.add_paragraph('defaultCurrency: "INR" (Indian Rupee)', style='List Bullet')
    doc.add_paragraph('defaultLocale: "en-IN"', style='List Bullet')
    
    doc.add_page_break()
    
    # 3.9 Setup & Deployment
    doc.add_heading('3.9 Setup & Deployment', level=2)
    
    doc.add_heading('Local Development Setup', level=3)
    
    setup_steps = [
        'Clone repository: `git clone <repo-url>`',
        'Navigate to project: `cd Bachatkaro`',
        'Install dependencies: `bun install` (or `npm install`)',
        'Create `.env.local` file with Supabase credentials',
        'Start dev server: `bun run dev` (or `npm run dev`)',
        'Open http://localhost:8080 in browser',
        'Make changes; Vite auto-reloads',
    ]
    
    for i, step in enumerate(setup_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    doc.add_heading('Production Build', level=3)
    doc.add_paragraph('Run: `bun run build` (or `npm run build`)', style='List Number')
    doc.add_paragraph('Creates optimized build in `dist/` folder', style='List Number')
    doc.add_paragraph('Deploy to Vercel: `vercel deploy` or via GitHub integration', style='List Number')
    
    doc.add_heading('Database Setup (Supabase)', level=3)
    doc.add_paragraph('Create new Supabase project', style='List Number')
    doc.add_paragraph('Run migrations: `supabase migration up`', style='List Number')
    doc.add_paragraph('Set up auth: Enable email/password auth in Supabase Dashboard', style='List Number')
    doc.add_paragraph('Configure RLS policies: All tables have Example policies (review in production)', style='List Number')
    
    doc.add_page_break()
    
    # 3.10 Testing Strategy
    doc.add_heading('3.10 Testing Strategy', level=2)
    
    doc.add_heading('Unit Testing', level=3)
    doc.add_paragraph('Framework: Vitest (Jest-compatible)', style='List Bullet')
    doc.add_paragraph('Examples: `src/test/example.test.ts`', style='List Bullet')
    doc.add_paragraph('Run tests: `bun run test`', style='List Bullet')
    doc.add_paragraph('Coverage: Aim for 70%+ coverage on utilities and calculations', style='List Bullet')
    
    doc.add_heading('Component Testing', level=3)
    doc.add_paragraph('Framework: Testing Library React + Vitest', style='List Bullet')
    doc.add_paragraph('Approach: Test user interactions, not implementation', style='List Bullet')
    doc.add_paragraph('Example: Test that "Add Expense" button calls submit handler', style='List Bullet')
    
    doc.add_heading('Manual QA', level=3)
    doc.add_paragraph('Test flows: Signup → Add Expense → Create Group → Split Bill', style='List Bullet')
    doc.add_paragraph('Voice input: Test in English and Hindi', style='List Bullet')
    doc.add_paragraph('Real-time sync: Open app on 2 devices, add expense, verify both update', style='List Bullet')
    doc.add_paragraph('Mobile: Test on actual mobile device (iOS Safari, Android Chrome)', style='List Bullet')
    doc.add_paragraph('Browser support: Chrome, Firefox, Safari, Edge', style='List Bullet')
    
    doc.add_heading('Performance Testing', level=3)
    doc.add_paragraph('Page load: Measure with Lighthouse', style='List Bullet')
    doc.add_paragraph('Target: > 90 Lighthouse score', style='List Bullet')
    doc.add_paragraph('Bundle size: Keep under 500KB gzipped', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== SECTION 4: MAINTENANCE & UPDATE GUIDE =====
    doc.add_heading('SECTION 4: MAINTENANCE & UPDATE GUIDE', level=1)
    
    doc.add_heading('4.1 How to Modify Existing Features', level=2)
    
    doc.add_heading('Example: Modify Expense Categories', level=3)
    doc.add_paragraph(
        'Current state: Categories are hardcoded in Dashboard.tsx'
    )
    doc.add_paragraph(
        'To change categories (e.g., add "Fitness" category):',
        style='List Bullet'
    )
    
    steps = [
        'Edit: src/pages/Dashboard.tsx → find category dropdown options',
        'Add new option: { label: "Fitness", value: "Fitness" }',
        'Update translation files: src/i18n/translations.ts (add Hindi label)',
        'Update voiceParser.ts: Add "gym", "fitness", "yoga" keywords → "Fitness"',
        'Test: Add expense with "500 for gym" via voice',
        'Deploy: Commit, push, Vercel auto-deploys',
    ]
    
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    doc.add_heading('Example: Change EMI Calculation Logic', level=3)
    doc.add_paragraph(
        'EMI calculations are in src/utils/loanCalculator.ts'
    )
    
    steps2 = [
        'Edit src/utils/loanCalculator.ts → function calculateReducingEMI()',
        'Modify calculation formula if needed (be careful with financial math!)',
        'Run tests: bun run test',
        'Test with known example: 5L principal, 9% rate, 60 months → should be ~9655 EMI',
        'Update Dashboard where getLoanSummary() is called',
        'Deploy once tested',
    ]
    
    for i, step in enumerate(steps2, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    
    doc.add_heading('Impact Analysis Checklist', level=3)
    doc.add_paragraph('Which files reference this feature?', style='List Bullet')
    doc.add_paragraph('Does database schema need changes (migrations)?', style='List Bullet')
    doc.add_paragraph('Do any tests break?', style='List Bullet')
    doc.add_paragraph('Does UI need adjustment?', style='List Bullet')
    doc.add_paragraph('Are there i18n strings to update?', style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('4.2 How to Add New Features', level=2)
    
    doc.add_heading('Example: Add "Budget Alert" Feature', level=3)
    doc.add_paragraph(
        'Goal: Notify user when they exceed monthly budget'
    )
    
    integration_steps = [
        "Step 1 — Database: Create `budget_alerts` table in migration file",
        "Step 2 — Component: Add BudgetAlert.tsx component in src/components/dashboard/",
        "Step 3 — Page: Import BudgetAlert into Dashboard.tsx, add to UI",
        "Step 4 — Logic: Query alerts via React Query, refetch on expense add",
        "Step 5 — Styling: Use shadcn Alert component for notification",
        "Step 6 — i18n: Add English/Hindi labels to translations.ts",
        "Step 7 — Test: Test manually, add unit tests",
        "Step 8 — Deploy: Commit, push, verify on Vercel",
    ]
    
    for step in integration_steps:
        doc.add_paragraph(step, style="List Bullet")
    
    doc.add_heading("Feature Integration Checklist", level=3)
    doc.add_paragraph("Create new TypeScript type (if needed)", style="List Bullet")
    doc.add_paragraph("Add database table/columns (write migration)", style="List Bullet")
    doc.add_paragraph("Create/modify React components", style="List Bullet")
    doc.add_paragraph("Add API calls (Supabase REST or RPC)", style="List Bullet")
    doc.add_paragraph("Integrate with state (Context/React Query)", style="List Bullet")
    doc.add_paragraph("Add i18n translations", style="List Bullet")
    doc.add_paragraph("Write unit + component tests", style="List Bullet")
    doc.add_paragraph("Test on mobile & desktop", style="List Bullet")
    doc.add_paragraph("Update documentation", style="List Bullet")
    
    doc.add_page_break()
    
    doc.add_heading('4.3 Risk Areas & Critical Components', level=2)
    
    doc.add_heading('High-Risk Components', level=3)
    
    risk_areas = [
        {
            'component': 'Loan Calculator (loanCalculator.ts)',
            'risk': 'CRITICAL - Financial calculations must be 100% accurate',
            'impact': 'Wrong EMI calculation loses user trust',
            'mitigation': 'Test against known examples, add unit tests, document formulas'
        },
        {
            'component': 'Bill Splitting (splitCalculator.ts)',
            'risk': 'HIGH - Complex logic for debt settlement',
            'impact': 'Wrong calculations cause disputes among users',
            'mitigation': 'Test with multiple scenarios, simplify debt algorithm'
        },
        {
            'component': 'Authentication (AuthContext.tsx)',
            'risk': 'CRITICAL - Security & user access',
            'impact': 'Compromised auth = full data exposure',
            'mitigation': 'Use Supabase Auth (battle-tested), never store tokens in localStorage'
        },
        {
            'component': 'Real-time Sync (Supabase Realtime)',
            'risk': 'MEDIUM - Race conditions under high load',
            'impact': 'Duplicate expenses or inconsistent balances',
            'mitigation': 'Use database constraints, add idempotency keys'
        },
        {
            'component': 'Voice Parser (voiceParser.ts)',
            'risk': 'MEDIUM - Can misinterpret user intent',
            'impact': 'Wrong expense logged (wrong amount or category)',
            'mitigation': 'Always show confirmation before saving, allow user edit'
        },
    ]
    
    for area in risk_areas:
        doc.add_paragraph(f"{area['component']}", style='List Number')
        doc.add_paragraph(f"Risk: {area['risk']}", style='List Bullet 2')
        doc.add_paragraph(f"Impact: {area['impact']}", style='List Bullet 2')
        doc.add_paragraph(f"Mitigation: {area['mitigation']}", style='List Bullet 2')
    
    doc.add_heading('Data Security Considerations', level=3)
    doc.add_paragraph('All user data is encrypted at rest in Supabase', style='List Bullet')
    doc.add_paragraph('TLS 1.2+ for all data in transit', style='List Bullet')
    doc.add_paragraph('Row-Level Security (RLS) prevents cross-user data access', style='List Bullet')
    doc.add_paragraph('Never log sensitive data (passwords, PII)', style='List Bullet')
    doc.add_paragraph('Regularly rotate Supabase API keys', style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('4.4 Performance Optimization Opportunities', level=2)
    
    optimizations = [
        {
            'area': 'Database Queries',
            'current': 'Some queries load all expenses, filter in frontend',
            'improvement': 'Add database-level filtering (WHERE clauses)',
            'effort': 'Low',
            'impact': 'Medium'
        },
        {
            'area': 'Charts (Recharts)',
            'current': 'Re-render on every prop change',
            'improvement': 'Memoize chart components, use React.memo()',
            'effort': 'Low',
            'impact': 'Medium'
        },
        {
            'area': 'Real-time Subscriptions',
            'current': 'Subscribe to all changes in a table',
            'improvement': 'Filter subscriptions (only user\'s data)',
            'effort': 'Medium',
            'impact': 'High'
        },
        {
            'area': 'Bundle Size',
            'current': 'Full chart library + all UI components',
            'improvement': 'Code-split Analytics (lazy load), tree-shake unused components',
            'effort': 'Medium',
            'impact': 'Low-Medium'
        },
        {
            'area': 'Voice Recognition',
            'current': 'Uses browser Web Speech API (can be slow)',
            'improvement': 'Consider serverless speech-to-text (Google Cloud Speech)',
            'effort': 'High',
            'impact': 'Medium'
        },
        {
            'area': 'Images/Icons',
            'current': 'Lucide React (great but many icons loaded)',
            'improvement': 'Cache SVGs, use WebP for any raster images',
            'effort': 'Low',
            'impact': 'Low'
        },
        {
            'area': 'Caching Strategy',
            'current': 'React Query cache expires after 2 minutes',
            'improvement': 'Optimize staleTime based on data volatility',
            'effort': 'Low',
            'impact': 'Medium'
        },
    ]
    
    for opt in optimizations:
        doc.add_paragraph(f"{opt['area']}", style='List Number')
        doc.add_paragraph(f"Current: {opt['current']}", style='List Bullet 2')
        doc.add_paragraph(f"Improvement: {opt['improvement']}", style='List Bullet 2')
        doc.add_paragraph(f"Effort: {opt['effort']}, Impact: {opt['impact']}", style='List Bullet 2')
    
    doc.add_page_break()
    
    # ===== SECTION 5: APPENDIX =====
    doc.add_heading('SECTION 5: APPENDIX', level=1)
    
    doc.add_heading('5.1 Glossary', level=2)
    
    glossary_terms = [
        ('Expense', 'A recorded transaction where money is spent'),
        ('Category', 'Classification of expense (Food, Travel, Shopping, Bills)'),
        ('EMI', 'Equated Monthly Installment — fixed monthly loan payment'),
        ('Bill Splitting', 'Dividing shared expense among multiple people'),
        ('Settlement', 'Marking debts as paid between group members'),
        ('Group', 'Set of users for shared expense management'),
        ('Voice Input', 'Hands-free expense entry via speech-to-text'),
        ('Reducing Balance', 'Loan interest calculation method (most common for mortgages)'),
        ('Flat Interest', 'Loan interest calculation method (used for some personal loans)'),
        ('RLS', 'Row-Level Security — database policy for data access control'),
        ('Supabase', 'Backend-as-a-Service platform with PostgreSQL database'),
        ('React Query', 'Library for managing server state and caching'),
        ('Vite', 'Modern build tool for JavaScript applications'),
        ('TypeScript', 'JavaScript superset that adds static typing'),
        ('Tailwind CSS', 'Utility-first CSS framework'),
        ('shadcn-ui', 'Pre-built React components using Tailwind & Radix UI'),
    ]
    
    for term, definition in glossary_terms:
        p = doc.add_paragraph(f'{term}: ', style='List Number')
        p_run = p.runs[0]
        p_run.bold = True
        p.add_run(definition)
    
    doc.add_page_break()
    
    doc.add_heading('5.2 Assumptions & Known Issues', level=2)
    
    doc.add_heading('Assumptions', level=3)
    assumptions = [
        'Users have stable internet connection (app requires online access)',
        'Browser supports Web Speech API (Chrome, Edge, Safari)',
        'Database scales to 100K+ concurrent users',
        'Supabase pricing remains feasible',
        'Users are primarily in India (INR currency)',
        'Mobile users will eventually migrate to native app',
    ]
    for assumption in assumptions:
        doc.add_paragraph(assumption, style='List Bullet')
    
    doc.add_heading('Known Issues', level=3)
    known_issues = [
        'Issue: Voice recognition sometimes mishears expense amount (e.g., "500" as "5000") → Mitigation: Show confirmation before saving',
        'Issue: Real-time sync can lag under high load → Mitigation: Manual refresh button available',
        'Issue: Trip planning destinationdata is hardcoded → Future: Integrate live API for dynamic data',
        'Issue: No offline mode currently → Planned for Phase 2',
        'Issue: Mobile UI can be cramped on very small screens (< 320px) → Tested on iPhone SE',
        'Issue: Dark mode not perfect on all components → Ongoing refinement',
    ]
    for issue in known_issues:
        doc.add_paragraph(issue, style='List Bullet')
    
    doc.add_heading('5.3 Open Questions & TODOs', level=2)
    
    doc.add_paragraph('Should we implement offline mode using Capacitor?', style='List Bullet')
    doc.add_paragraph('What is the timeline for native mobile apps?', style='List Bullet')
    doc.add_paragraph('Should we support other loan interest calculation methods?', style='List Bullet')
    doc.add_paragraph('How do we prevent fraud in group bill splitting?', style='List Bullet')
    doc.add_paragraph('Should we add AI chatbot for financial advice?', style='List Bullet')
    doc.add_paragraph('What is our multi-currency strategy for global expansion?', style='List Bullet')
    doc.add_paragraph('How do we monetize (freemium model, premium features)?', style='List Bullet')
    doc.add_paragraph('Should we integrate with bank APIs for auto-sync?', style='List Bullet')
    
    doc.add_page_break()
    
    # ===== FINAL PAGE =====
    doc.add_heading('Document Information', level=1)
    
    doc.add_paragraph('Document Version: 1.0', style='List Bullet')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M UTC")}', style='List Bullet')
    doc.add_paragraph('Framework: React 18 + TypeScript + Vite', style='List Bullet')
    doc.add_paragraph('Backend: Supabase (PostgreSQL) + Vercel', style='List Bullet')
    doc.add_paragraph('Status: Beta / Production Ready', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('How to Update This Documentation', level=2)
    
    doc.add_paragraph('This document should be regenerated whenever:', style='List Bullet')
    doc.add_paragraph('Major features are added or deprecated', style='List Bullet 2')
    doc.add_paragraph('Tech stack changes significantly', style='List Bullet 2')
    doc.add_paragraph('Database schema is updated', style='List Bullet 2')
    doc.add_paragraph('Architecture decisions are revisited', style='List Bullet 2')
    doc.add_paragraph('New team members onboard (quarterly update)', style='List Bullet 2')
    
    # Save document
    doc_path = 'd:/MyProject/Bachatkaro/BACHAT_KARO_COMPREHENSIVE_DOCUMENTATION.docx'
    doc.save(doc_path)
    print(f'✅ Documentation generated successfully!')
    print(f'📄 File: {doc_path}')
    print(f'📊 Sections: 5 (Executive Summary, PRD, Technical Doc, Maintenance, Appendix)')
    print(f'📝 Pages: ~50+ pages of professional documentation')

if __name__ == '__main__':
    generate_documentation()
