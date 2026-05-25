#!/usr/bin/env python3
"""
Convert BachatKaro Markdown Documentation to Word Documents
Creates professional .docx files from markdown documentation
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level=1, color=(41, 41, 41)):
    """Add heading with custom color"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def parse_markdown_and_create_docx(md_file, output_file):
    """
    Parse markdown file and create Word document
    """
    print(f"📄 Converting: {md_file}")
    
    if not os.path.exists(md_file):
        print(f"  ❌ File not found: {md_file}")
        return False
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Add title page styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Parse and add content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines unless they're meaningful
        if not line.strip():
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            add_heading_with_color(doc, line.replace('# ', '').strip(), level=1, color=(0, 102, 204))
            i += 1
        elif line.startswith('## '):
            add_heading_with_color(doc, line.replace('## ', '').strip(), level=2, color=(51, 51, 51))
            i += 1
        elif line.startswith('### '):
            add_heading_with_color(doc, line.replace('### ', '').strip(), level=3, color=(102, 102, 102))
            i += 1
        elif line.startswith('#### '):
            add_heading_with_color(doc, line.replace('#### ', '').strip(), level=4, color=(128, 128, 128))
            i += 1
        elif line.startswith('##### '):
            doc.add_heading(line.replace('##### ', '').strip(), level=5)
            i += 1
        
        # Bold text (before paragraphs to avoid issues)\n        elif line.strip().startswith('**') and line.strip().endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', '').strip())
            run.bold = True
            i += 1
        
        # Code blocks
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            # Detect language (optional)
            if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('```'):
                # Could be language identifier
                if lines[i].strip() in ['python', 'typescript', 'tsx', 'bash', 'sql', 'json', 'bash', 'javascript']:
                    i += 1
            
            # Collect code block
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block
            if code_lines:
                code_text = '\n'.join(code_lines).rstrip()
                p = doc.add_paragraph(code_text, style='List Number')
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(51, 51, 51)
            
            i += 1  # Skip closing ```
        
        # Tables (markdown format)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Parse table
            header_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            i += 1  # Skip separator line
            i += 1
            
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row_cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if len(row_cells) == len(header_cells):
                    rows.append(row_cells)
                i += 1
            
            # Create table
            if header_cells and rows:
                table = doc.add_table(rows=1, cols=len(header_cells))
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_row = table.rows[0]
                for idx, cell_text in enumerate(header_cells):
                    cell = header_row.cells[idx]
                    cell.text = cell_text
                    set_cell_background(cell, 'E7E6E6')
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                
                # Data rows
                for row_data in rows:
                    row = table.add_row()
                    for idx, cell_text in enumerate(row_data):
                        row.cells[idx].text = cell_text.replace('`', '')
        
        # Lists
        elif line.strip().startswith('- '):
            text = line.replace('- ', '').strip()
            doc.add_paragraph(text, style='List Bullet')
            i += 1
        elif line.strip().startswith('* '):
            text = line.replace('* ', '').strip()
            doc.add_paragraph(text, style='List Bullet')
            i += 1
        elif line.strip() and line.strip()[0].isdigit() and '. ' in line:
            text = re.sub(r'^\d+\.\s+', '', line).strip()
            doc.add_paragraph(text, style='List Number')
            i += 1
        
        # Regular paragraphs
        elif line.strip():
            # Replace markdown formatting
            text = line.strip()
            # Replace **text** with just text (will be handled separately)
            text = text.replace('**', '')
            text = text.replace('`', '')
            text = text.replace('[', '').replace(']', '')
            text = re.sub(r'\(http[s]?://[^\)]+\)', '', text)
            
            if text.strip():
                doc.add_paragraph(text)
            i += 1
        else:
            i += 1
    
    # Save document
    doc.save(output_file)
    print(f"  ✅ Created: {output_file}")
    return True

def main():
    """Main function to convert all documentation files"""
    print("\n" + "="*70)
    print("🚀 BachatKaro Documentation to Word Converter")
    print("="*70 + "\n")
    
    # Documentation files to convert
    doc_files = [
        ('DEVELOPER_README.md', 'DEVELOPER_README.docx'),
        ('SETUP_GUIDE.md', 'SETUP_GUIDE.docx'),
        ('ARCHITECTURE.md', 'ARCHITECTURE.docx'),
        ('API_DOCUMENTATION.md', 'API_DOCUMENTATION.docx'),
        ('DATABASE.md', 'DATABASE.docx'),
        ('COMPONENTS.md', 'COMPONENTS.docx'),
        ('DEPLOYMENT.md', 'DEPLOYMENT.docx'),
        ('CONTRIBUTING.md', 'CONTRIBUTING.docx'),
        ('TROUBLESHOOTING.md', 'TROUBLESHOOTING.docx'),
        ('DOCUMENTATION_COMPLETE.md', 'DOCUMENTATION_COMPLETE.docx'),
        ('USER_DOCUMENTATION.md', 'USER_DOCUMENTATION.docx'),
        ('QUICK_START_GUIDE.md', 'QUICK_START_GUIDE.docx'),
        ('STAKEHOLDER_GUIDE.md', 'STAKEHOLDER_GUIDE.docx'),
    ]
    
    success_count = 0
    failed_count = 0
    
    for md_file, output_file in doc_files:
        md_path = os.path.join('.', md_file)
        output_path = os.path.join('.', output_file)
        
        if parse_markdown_and_create_docx(md_path, output_path):
            success_count += 1
        else:
            failed_count += 1
    
    print("\n" + "="*70)
    print(f"✅ Conversion Complete!")
    print(f"   Created: {success_count} Word documents")
    print(f"   Failed: {failed_count}")
    print("="*70)
    print("\n📁 All .docx files are ready in the project root directory!")
    print("📧 You can now use these for creating ads and marketing materials!\n")

if __name__ == '__main__':
    main()
